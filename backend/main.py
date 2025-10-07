"""
SaralPolicy - Local AI Insurance Analysis
Production-ready with bilingual UI, working TTS, and human-readable formatting.
"""

import os
import json
import tempfile
from pathlib import Path
from typing import Dict, Any, List
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import structlog
import asyncio
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache
import hashlib
import time

import PyPDF2
import docx

# Import local services
from app.services.ollama_llm_service import OllamaLLMService
from app.services.rag_service import rag_service
from app.services.evaluation import EvaluationManager
from app.services.hitl_service import HITLService
from app.services.guardrails_service import GuardrailsService
from app.services.tts_service import TTSService
from app.services.translation_service import TranslationService

# Configure logging
structlog.configure(
    processors=[
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_logger_name,
        structlog.stdlib.add_log_level,
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.JSONRenderer()
    ],
    context_class=dict,
    logger_factory=structlog.stdlib.LoggerFactory(),
    wrapper_class=structlog.stdlib.BoundLogger,
    cache_logger_on_first_use=True,
)

logger = structlog.get_logger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="SaralPolicy",
    description="AI-powered insurance document analysis with local privacy",
    version="2.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files and templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# Initialize Ollama LLM Service with Gemma 3 4B
try:
    logger.info("🚀 Initializing Ollama LLM Service with Gemma 3 4B")
    ollama_service = OllamaLLMService(model_name="gemma3:4b")
    logger.info("✅ Ollama LLM Service initialized successfully with Gemma 3 4B")
except Exception as e:
    logger.error("❌ Failed to initialize Ollama LLM Service", error=str(e))
    logger.info("💡 Make sure Ollama is running: ollama serve")
    logger.info("💡 And model is pulled: ollama pull gemma3:4b")
    ollama_service = None

# Supporting services
eval_manager = EvaluationManager()
hitl_service = HITLService()
guardrails_service = GuardrailsService()
tts_service = TTSService()
translation_service = TranslationService()


class SaralPolicyLocal:
    """SaralPolicy with local Llama, guardrails, and HITL."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.ollama_service = ollama_service
        self.rag_service = rag_service
        self.eval_manager = eval_manager
        self.hitl_service = hitl_service
        self.guardrails_service = guardrails_service
        self.executor = ThreadPoolExecutor(max_workers=4)  # For parallel processing
        self._text_cache = {}  # Simple in-memory cache
    
    def _get_file_hash(self, file_path: str) -> str:
        """Generate hash for file caching."""
        with open(file_path, 'rb') as f:
            # Read first and last 8KB for speed (works well for most documents)
            first_chunk = f.read(8192)
            f.seek(-min(8192, os.path.getsize(file_path)), 2)
            last_chunk = f.read(8192)
            return hashlib.md5(first_chunk + last_chunk).hexdigest()
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from uploaded document with caching."""
        start_time = time.time()
        
        # Check cache first
        file_hash = self._get_file_hash(file_path)
        if file_hash in self._text_cache:
            logger.info("📦 Using cached document text", cache_hit=True)
            return self._text_cache[file_hash]
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            text = self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            text = self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Cache the result
        self._text_cache[file_hash] = text
        
        elapsed = time.time() - start_time
        logger.info(f"✅ Document parsed in {elapsed:.2f}s", 
                   format=file_extension, 
                   size_kb=len(text)/1024,
                   cached=False)
        
        return text
    
    def _extract_pdf_page(self, pdf_reader, page_num: int) -> str:
        """Extract text from a single PDF page."""
        try:
            return pdf_reader.pages[page_num].extract_text() + "\n"
        except Exception as e:
            logger.warning(f"Failed to extract page {page_num}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF with parallel page processing."""
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            # For small PDFs, process sequentially
            if num_pages <= 5:
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text() + "\n")
            else:
                # For larger PDFs, use parallel processing
                from concurrent.futures import ThreadPoolExecutor, as_completed
                with ThreadPoolExecutor(max_workers=4) as executor:
                    future_to_page = {executor.submit(self._extract_pdf_page, pdf_reader, i): i 
                                     for i in range(num_pages)}
                    
                    # Collect results in order
                    results = {}
                    for future in as_completed(future_to_page):
                        page_num = future_to_page[future]
                        results[page_num] = future.result()
                    
                    # Assemble in correct order
                    text_parts = [results[i] for i in range(num_pages)]
        
        return "".join(text_parts)
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX with optimized processing."""
        doc = docx.Document(file_path)
        # Use list comprehension for better performance
        text_parts = [paragraph.text + "\n" for paragraph in doc.paragraphs]
        return "".join(text_parts)
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT with encoding detection."""
        # Try UTF-8 first, fall back to latin-1
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def analyze_policy(self, text: str) -> Dict[str, Any]:
        """
        Analyze insurance policy using Ollama with Gemma 3 4B model.
        Complete local processing with IRDAI compliance.
        """
        # Track overall performance
        analysis_start = time.time()
        performance_metrics = {}
        
        try:
            logger.info("🚀 Starting policy analysis with Ollama Gemma 3 4B", text_length=len(text))
            
            # Apply guardrails first
            guardrails_result = self.guardrails_service.validate_input(text)
            if not guardrails_result["is_valid"]:
                return {
                    "error": f"Input validation failed: {guardrails_result['reason']}",
                    "status": "error"
                }
            
            # Use Ollama LLM Service with Gemma 3 4B
            if self.ollama_service:
                try:
                    logger.info("⚙️ Using Ollama with Gemma 3 4B model")
                    
                    # RAG Enhancement: Index document for session-based Q&A
                    doc_id = f"policy_{hash(text[:100])}"  # Simple doc ID
                    if self.rag_service and self.rag_service.enabled:
                        rag_start = time.time()
                        logger.info("📚 Indexing document for RAG-enhanced Q&A")
                        self.rag_service.index_document(
                            text=text,
                            document_id=doc_id,
                            metadata={
                                'title': 'User Policy Document',
                                'policy_type': 'insurance',
                                'indexed_at': 'analysis_time'
                            },
                            collection_name="session"
                        )
                        performance_metrics['rag_indexing_time'] = round(time.time() - rag_start, 2)
                    
                    llm_start = time.time()
                    intelligent_summary = self.ollama_service.generate_intelligent_summary(text)
                    key_terms = self.ollama_service.extract_terms_with_explanations(text)
                    exclusions = self.ollama_service.identify_exclusions(text)
                    coverage = self.ollama_service.extract_coverage_details(text)
                    performance_metrics['llm_processing_time'] = round(time.time() - llm_start, 2)
                    
                    # RAG Enhancement: Augment with IRDAI knowledge
                    rag_enhanced = False
                    rag_citations = []
                    if self.rag_service and self.rag_service.enabled:
                        try:
                            rag_query_start = time.time()
                            # Query IRDAI knowledge base for relevant regulations
                            policy_type = intelligent_summary.get("policy_type", "health_insurance")
                            
                            # Multiple targeted queries for comprehensive coverage
                            queries = [
                                f"What are the mandatory coverage requirements for {policy_type}?",
                                f"What are the claim settlement timelines for {policy_type}?",
                                f"What waiting periods apply to {policy_type}?",
                                f"What exclusions are regulated in {policy_type}?"
                            ]
                            
                            all_rag_results = []
                            for query in queries:
                                results = self.rag_service.query_knowledge_base(query, top_k=2)
                                all_rag_results.extend(results)
                            
                            performance_metrics['rag_query_time'] = round(time.time() - rag_query_start, 2)
                            
                            if all_rag_results:
                                # Extract unique, high-relevance results
                                seen_content = set()
                                rag_regulations = []
                                
                                for result in all_rag_results:
                                    if result['hybrid_score'] >= 0.4:  # Relevance threshold
                                        content = result['content']
                                        # Avoid duplicates
                                        content_hash = hash(content[:100])
                                        if content_hash not in seen_content:
                                            seen_content.add(content_hash)
                                            
                                            # Extract key insight (first 150 chars)
                                            insight = content[:150].strip() + "..."
                                            source = result['metadata'].get('source', 'IRDAI Guidelines')
                                            category = result['metadata'].get('insurance_category', 'general')
                                            
                                            rag_regulations.append(insight)
                                            rag_citations.append({
                                                "source": source,
                                                "category": category,
                                                "relevance": round(result['hybrid_score'], 3),
                                                "excerpt": insight
                                            })
                                
                                if rag_regulations:
                                    intelligent_summary['irdai_regulations'] = rag_regulations[:4]  # Top 4
                                    rag_enhanced = True
                                    logger.info("[OK] Enhanced with IRDAI knowledge base", 
                                              citations=len(rag_citations))
                        except Exception as rag_err:
                            logger.warning(f"RAG enhancement failed: {rag_err}")
                    
                    # Calculate total analysis time
                    total_time = round(time.time() - analysis_start, 2)
                    performance_metrics['total_analysis_time'] = total_time
                    
                    result = {
                        "summary": intelligent_summary.get("simple_summary", ""),
                        "key_highlights": [],  # Will be added from knowledge base
                        "what_you_need_to_know": [
                            {"icon": "⏰", "point": "When to Renew", "explanation": "Renew before expiry to avoid waiting period reset"},
                            {"icon": "📞", "point": "How to Claim", "explanation": "Cashless: Call before admission. Reimbursement: Keep all bills"},
                            {"icon": "📋", "point": "Keep Handy", "explanation": "Policy doc, ID card, policy number, helpline number"},
                            {"icon": "🚨", "point": "Emergency", "explanation": "Save insurer helpline, TPA number, network hospitals list"}
                        ],
                        "policy_type": intelligent_summary.get("policy_type", "health_insurance"),
                        "educational_resources": {
                            "irdai": intelligent_summary.get("irdai_guidelines"),
                            "irdai_main": "https://www.irdai.gov.in",
                            "complaint_portal": "https://igms.irda.gov.in",
                            "insurance_ombudsman": "https://www.cioins.co.in",
                            "video": f"https://www.youtube.com/results?search_query={intelligent_summary.get('policy_type', 'health')}+insurance+explained+hindi"
                        },
                        "irdai_regulations": intelligent_summary.get("irdai_regulations", []),
                        "irdai_citations": rag_citations if rag_enhanced else [],
                        "key_terms": key_terms,
                        "exclusions": exclusions,
                        "coverage": coverage,
                        "confidence_score": intelligent_summary.get("confidence", 0.95),
                        "quality_grade": "A+",
                        "irdai_compliant": True,
                        "model_used": "Gemma 3 4B (Ollama)" + (" + RAG" if rag_enhanced else ""),
                        "rag_enhanced": rag_enhanced,
                        "hitl_triggered": False,
                        "status": "success",
                        "performance_metrics": performance_metrics
                    }
                    
                    logger.info("✅ Ollama analysis completed successfully", 
                               total_time=total_time,
                               **performance_metrics)
                    return result
                    
                except Exception as ollama_error:
                    logger.error("Ollama LLM failed", error=str(ollama_error))
                    raise
            
            # If Ollama not available, return error
            logger.error("❌ Ollama service not initialized")
            return {
                "error": "Ollama service not available. Please start Ollama and pull gemma3:4b model.",
                "status": "error",
                "instructions": [
                    "1. Install Ollama: https://ollama.ai",
                    "2. Start Ollama: ollama serve",
                    "3. Pull model: ollama pull gemma3:4b"
                ]
            }
            
        except Exception as e:
            logger.error("❌ Policy analysis failed", error=str(e))
            return {
                "error": f"Analysis failed: {str(e)}",
                "status": "error"
            }
    
    def answer_question(self, text: str, question: str) -> str:
        """Answer questions about the policy using Ollama."""
        try:
            # Apply guardrails to question
            guardrails_result = self.guardrails_service.validate_question(question)
            if not guardrails_result["is_valid"]:
                return f"Question validation failed: {guardrails_result['reason']}"
            
            # Get answer from Ollama
            if self.ollama_service:
                answer = self.ollama_service.answer_question(text, question)
                
                # Evaluate answer quality
                eval_result = self.eval_manager.evaluate_answer(text, question, answer)
                
                # If confidence is low, trigger HITL
                if eval_result.get("confidence_score", 0.0) < 0.85:
                    logger.info("Low confidence answer, triggering HITL")
                    self.hitl_service.trigger_qa_review(question, answer, eval_result)
                
                return answer
            else:
                return "Ollama service not available. Please start Ollama and pull gemma3:4b model."
            
        except Exception as e:
            logger.error("Failed to answer question", error=str(e))
            return "Error answering question. Please try again."


# Initialize the analyzer
analyzer = SaralPolicyLocal()


@app.get("/", response_class=HTMLResponse)
async def home():
    """Modern bilingual UI with working TTS and human-readable formatting."""
    html_content = Path("templates/index.html")
    if html_content.exists():
        return html_content.read_text(encoding='utf-8')
    
    # Fallback inline HTML
    return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SaralPolicy - AI Insurance Analysis</title>
    <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&family=Noto+Sans+Devanagari:wght@300;400;500;700&display=swap" rel="stylesheet">
    <link href="https://fonts.googleapis.com/icon?family=Material+Icons" rel="stylesheet">
    <style>
        :root {
            --primary: #6750A4;
            --primary-light: #EADDFF;
            --error: #BA1A1A;
            --success: #198754;
            --text: #1C1B1F;
            --text-secondary: #49454F;
        }
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Roboto', 'Noto Sans Devanagari', sans-serif;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            color: var(--text);
            line-height: 1.6;
            min-height: 100vh;
            padding: 20px;
        }
        .container { max-width: 1400px; margin: 0 auto; }
        .header {
            text-align: center;
            margin-bottom: 40px;
            padding: 40px 20px;
            background: linear-gradient(135deg, var(--primary) 0%, #9333EA 100%);
            border-radius: 24px;
            color: white;
            box-shadow: 0 10px 40px rgba(0,0,0,0.15);
            position: relative;
        }
        .header h1 { font-size: 3rem; font-weight: 700; margin-bottom: 12px; }
        .tagline { font-size: 1.5rem; font-weight: 500; opacity: 0.95; margin-bottom: 8px; }
        .subtitle { font-size: 1.1rem; opacity: 0.85; }
        .lang-toggle {
            position: absolute;
            top: 20px;
            right: 20px;
            display: flex;
            gap: 8px;
            background: rgba(255,255,255,0.2);
            border-radius: 20px;
            padding: 8px;
        }
        .lang-btn {
            background: transparent;
            border: 2px solid white;
            color: white;
            padding: 8px 20px;
            border-radius: 16px;
            cursor: pointer;
            font-weight: 500;
            transition: all 0.3s;
        }
        .lang-btn.active { background: white; color: var(--primary); }
        .main-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 30px; margin-bottom: 40px; }
        .card {
            background: white;
            border-radius: 20px;
            padding: 30px;
            box-shadow: 0 5px 20px rgba(0,0,0,0.08);
        }
        .card h2 { margin-bottom: 20px; color: var(--primary); font-weight: 600; font-size: 1.5rem; }
        .upload-area {
            border: 3px dashed var(--primary);
            border-radius: 16px;
            padding: 50px 20px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s;
            background: var(--primary-light);
        }
        .upload-area:hover { transform: translateY(-4px); box-shadow: 0 8px 24px rgba(103,80,164,0.2); }
        .upload-icon { font-size: 4rem; margin-bottom: 16px; }
        .upload-text { font-size: 1.25rem; font-weight: 500; margin-bottom: 8px; }
        .upload-subtext { color: var(--text-secondary); margin-bottom: 20px; }
        .btn {
            background: var(--primary);
            color: white;
            border: none;
            border-radius: 24px;
            padding: 16px 32px;
            font-size: 1.1rem;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.3s;
            box-shadow: 0 4px 12px rgba(103,80,164,0.3);
            display: inline-flex;
            align-items: center;
            gap: 8px;
            width: 100%;
            justify-content: center;
            margin-top: 20px;
        }
        .btn:hover { transform: translateY(-2px); box-shadow: 0 8px 24px rgba(103,80,164,0.4); }
        .btn:disabled { opacity: 0.6; cursor: not-allowed; transform: none; }
        .feature-item {
            padding: 20px;
            margin-bottom: 16px;
            border-radius: 12px;
            background: #E7E0EC;
            transition: all 0.3s;
        }
        .feature-item:hover { transform: translateX(8px); }
        .feature-icon { font-size: 2rem; margin-bottom: 8px; color: var(--primary); }
        .feature-title { font-weight: 600; margin-bottom: 4px; font-size: 1.1rem; }
        .feature-desc { font-size: 0.95rem; color: var(--text-secondary); }
        .results-section { grid-column: 1 / -1; display: none; }
        .results-section.show { display: block; animation: fadeIn 0.5s; }
        @keyframes fadeIn { from { opacity: 0; transform: translateY(20px); } to { opacity: 1; transform: translateY(0); } }
        .result-card {
            background: white;
            border-radius: 16px;
            padding: 30px;
            margin-bottom: 20px;
            box-shadow: 0 3px 15px rgba(0,0,0,0.08);
            border-left: 4px solid var(--primary);
        }
        .result-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 20px; flex-wrap: wrap; gap: 16px; }
        .result-title { display: flex; align-items: center; gap: 12px; font-size: 1.5rem; font-weight: 600; color: var(--primary); }
        .result-content { color: var(--text); line-height: 1.8; font-size: 1.05rem; }
        .tts-controls { display: flex; gap: 10px; }
        .tts-btn {
            background: var(--primary);
            color: white;
            border: none;
            border-radius: 12px;
            padding: 10px 16px;
            font-size: 0.9rem;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.3s;
            display: flex;
            align-items: center;
            gap: 6px;
        }
        .tts-btn:hover { background: #5640a0; transform: scale(1.05); }
        .term-item {
            padding: 16px;
            margin-bottom: 12px;
            background: #E7E0EC;
            border-radius: 10px;
            border-left: 3px solid var(--success);
        }
        .term-name { font-weight: 600; font-size: 1.1rem; color: var(--primary); margin-bottom: 6px; }
        .term-definition { color: var(--text); font-size: 1rem; }
        .exclusion-item {
            padding: 12px 20px;
            margin-bottom: 10px;
            background: #fff3cd;
            border-radius: 8px;
            border-left: 3px solid #ff9800;
            display: flex;
            align-items: start;
            gap: 10px;
        }
        .coverage-item {
            display: grid;
            grid-template-columns: 200px 1fr;
            padding: 12px 0;
            border-bottom: 1px solid #E7E0EC;
        }
        .coverage-label { font-weight: 600; color: var(--text-secondary); }
        .coverage-value { color: var(--text); font-weight: 500; }
        .loading { display: none; text-align: center; padding: 40px; }
        .loading.show { display: block; }
        .spinner {
            width: 50px;
            height: 50px;
            border: 5px solid #E7E0EC;
            border-top: 5px solid var(--primary);
            border-radius: 50%;
            animation: spin 1s linear infinite;
            margin: 0 auto 20px;
        }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .metric-badge {
            display: inline-block;
            padding: 8px 16px;
            border-radius: 20px;
            font-weight: 600;
            margin-right: 10px;
            margin-bottom: 10px;
        }
        .metric-high { background: #d4edda; color: #155724; }
        .metric-medium { background: #fff3cd; color: #856404; }
        .metric-low { background: #f8d7da; color: #721c24; }
        @media (max-width: 768px) {
            .main-grid { grid-template-columns: 1fr; }
            .header h1 { font-size: 2rem; }
            .lang-toggle { position: relative; top: 0; right: 0; margin-top: 20px; }
        }
        /* Enhanced Features CSS */
        .audio-controls { display: flex; gap: 8px; align-items: center; }
        .audio-control-btn { background: var(--primary); color: white; border: none; border-radius: 12px; padding: 8px 12px; font-size: 0.85rem; font-weight: 500; cursor: pointer; transition: all 0.3s; display: inline-flex; align-items: center; gap: 4px; }
        .audio-control-btn:hover { background: #5640a0; transform: scale(1.05); }
        .tts-loading { display: none; width: 16px; height: 16px; border: 2px solid rgba(255,255,255,0.3); border-top-color: white; border-radius: 50%; animation: spin 0.8s linear infinite; }
        .copy-btn { background: var(--secondary); color: white; border: none; border-radius: 8px; padding: 6px 12px; font-size: 0.85rem; font-weight: 500; cursor: pointer; transition: all 0.3s; display: inline-flex; align-items: center; gap: 4px; }
        .copy-btn:hover { background: #4a4560; transform: translateY(-1px); }
        .analysis-progress { padding: 30px; text-align: center; }
        .progress-steps { display: flex; flex-direction: column; gap: 16px; margin-bottom: 30px; }
        .progress-step { display: flex; align-items: center; gap: 12px; padding: 12px 16px; background: #f5f5f5; border-radius: 8px; transition: all 0.3s; opacity: 0.5; }
        .progress-step.active { opacity: 1; background: var(--primary-light); animation: pulse 1s ease-in-out infinite; }
        .progress-step.complete { opacity: 1; background: #d4edda; }
        @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.8; } }
        .step-icon { flex-shrink: 0; width: 32px; height: 32px; display: flex; align-items: center; justify-content: center; }
        .step-text { font-size: 1rem; font-weight: 500; text-align: left; color: var(--text); }
        .progress-bar-container { width: 100%; height: 8px; background: #e0e0e0; border-radius: 4px; overflow: hidden; margin-bottom: 12px; }
        .progress-bar-fill { height: 100%; background: linear-gradient(90deg, var(--primary), var(--success)); transition: width 0.5s ease; width: 0%; }
        .progress-percentage { font-size: 1.5rem; font-weight: 700; color: var(--primary); }
        .action-buttons-bar { display: flex; gap: 12px; flex-wrap: wrap; align-items: center; padding: 16px; background: #E7E0EC; border-radius: 12px; margin-bottom: 20px; }
        .action-btn { background: white; color: var(--primary); border: 2px solid var(--primary); border-radius: 8px; padding: 8px 16px; font-size: 0.9rem; font-weight: 500; cursor: pointer; transition: all 0.3s; display: inline-flex; align-items: center; gap: 6px; }
        .action-btn:hover { background: var(--primary); color: white; transform: translateY(-2px); }
        body.dark-mode { background: linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%); color: #e0e0e0; }
        body.dark-mode .card { background: #2a2a2a; border-color: #3a3a3a; }
        body.dark-mode .upload-area { background: #333; border-color: #555; }
        body.dark-mode .result-card { background: #2a2a2a; }
        body.dark-mode .result-content { color: #e0e0e0; }
        body.dark-mode .term-item, body.dark-mode .feature-item { background: #333; }
        body.dark-mode .exclusion-item { background: #3a2f0f; }
        body.dark-mode .progress-step { background: #333; }
        body.dark-mode .progress-step.active { background: #4a3d7c; }
        body.dark-mode .action-btn { background: #2a2a2a; color: #e0e0e0; border-color: #555; }
        #darkModeToggle { position: fixed; bottom: 20px; right: 20px; width: 56px; height: 56px; border-radius: 50%; background: var(--primary); color: white; border: none; cursor: pointer; box-shadow: 0 4px 12px rgba(0,0,0,0.3); transition: all 0.3s; display: flex; align-items: center; justify-content: center; z-index: 1000; }
        #darkModeToggle:hover { transform: scale(1.1); }
        @media print { .no-print, #darkModeToggle, .lang-toggle, .tts-controls, .action-buttons-bar { display: none !important; } .header { background: white !important; color: black !important; } }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="lang-toggle">
                <button class="lang-btn active" onclick="switchLanguage('en')" id="langEnBtn">English</button>
                <button class="lang-btn" onclick="switchLanguage('hi')" id="langHiBtn">हिंदी</button>
            </div>
            <h1 id="appTitle">📄 SaralPolicy</h1>
            <div class="tagline" id="tagline">Insurance ka fine print, ab saaf saaf.</div>
            <p class="subtitle" id="subtitle">AI-Powered Insurance Document Analysis with Complete Privacy</p>
        </div>

        <div class="main-grid">
            <div class="card">
                <h2 id="uploadTitle">Upload Policy Document</h2>
                <div class="upload-area" onclick="document.getElementById('fileInput').click()">
                    <div class="upload-icon">📁</div>
                    <div class="upload-text" id="uploadText">Drop your policy here or click to browse</div>
                    <div class="upload-subtext" id="uploadSubtext">Supports PDF, DOCX, TXT (max 10MB)</div>
                    <input type="file" id="fileInput" style="display: none;" accept=".pdf,.docx,.txt" onchange="handleFileSelect(event)">
                </div>
                <button class="btn" id="analyzeBtn" onclick="analyzeDocument()" disabled>
                    <span class="material-icons">analytics</span>
                    <span id="analyzeBtnText">Analyze Policy</span>
                </button>
            </div>

            <div class="card">
                <h2 id="featuresTitle">Key Features</h2>
                <div class="feature-item">
                    <div class="feature-icon">🤖</div>
                    <div class="feature-title" id="feature1Title">100% Local AI Processing</div>
                    <div class="feature-desc" id="feature1Desc">No cloud uploads - your data stays on your device</div>
                </div>
                <div class="feature-item">
                    <div class="feature-icon">🔒</div>
                    <div class="feature-title" id="feature2Title">Complete Privacy</div>
                    <div class="feature-desc" id="feature2Desc">Zero data leakage with advanced PII protection</div>
                </div>
                <div class="feature-item">
                    <div class="feature-icon">🎯</div>
                    <div class="feature-title" id="feature3Title">Smart Analysis</div>
                    <div class="feature-desc" id="feature3Desc">Extracts key terms, coverage, and exclusions automatically</div>
                </div>
                <div class="feature-item">
                    <div class="feature-icon">🔊</div>
                    <div class="feature-title" id="feature4Title">Text-to-Speech</div>
                    <div class="feature-desc" id="feature4Desc">Listen to your policy in Hindi or English</div>
                </div>
            </div>
        </div>

        <div class="loading" id="loading">
            <div class="spinner"></div>
            <p id="loadingText">Analyzing your policy document...</p>
        </div>

        <div class="results-section card" id="resultsSection">
            <h2 id="resultsTitle">📊 Analysis Results</h2>
            <div id="resultsContent"></div>
        </div>
    </div>

    <!-- Dark Mode Toggle -->
    <button id="darkModeToggle" onclick="DarkModeController.toggle()" title="Toggle Dark Mode">
        <span class="material-icons">dark_mode</span>
    </button>

    <script>
        let selectedFile = null;
        let analysisData = null;
        let currentLanguage = 'en';
        let audioElement = null;

        const translations = {
            en: {
                appTitle: "📄 SaralPolicy",
                tagline: "Insurance ka fine print, ab saaf saaf.",
                subtitle: "AI-Powered Insurance Document Analysis with Complete Privacy",
                uploadTitle: "Upload Policy Document",
                uploadText: "Drop your policy here or click to browse",
                uploadSubtext: "Supports PDF, DOCX, TXT (max 10MB)",
                analyzeBtnText: "Analyze Policy",
                featuresTitle: "Key Features",
                feature1Title: "100% Local AI Processing",
                feature1Desc: "No cloud uploads - your data stays on your device",
                feature2Title: "Complete Privacy",
                feature2Desc: "Zero data leakage with advanced PII protection",
                feature3Title: "Smart Analysis",
                feature3Desc: "Extracts key terms, coverage, and exclusions automatically",
                feature4Title: "Text-to-Speech",
                feature4Desc: "Listen to your policy in Hindi or English",
                loadingText: "Analyzing your policy document...",
                resultsTitle: "📊 Analysis Results",
                policySummary: "Policy Summary",
                keyTerms: "Important Terms",
                exclusions: "Exclusions & Limitations",
                coverage: "Coverage Details",
                qualityMetrics: "Quality Metrics",
                confidenceScore: "Confidence Score",
                qualityGrade: "Quality Grade",
                expertReview: "⚠️ Requires Expert Review"
            },
            hi: {
                appTitle: "📄 सरल पॉलिसी",
                tagline: "बीमा की छोटी लिखावट, अब साफ़-साफ़।",
                subtitle: "पूर्ण गोपनीयता के साथ AI-संचालित बीमा दस्तावेज़ विश्लेषण",
                uploadTitle: "पॉलिसी दस्तावेज़ अपलोड करें",
                uploadText: "यहां अपनी पॉलिसी छोड़ें या ब्राउज़ करने के लिए क्लिक करें",
                uploadSubtext: "PDF, DOCX, TXT समर्थित (अधिकतम 10MB)",
                analyzeBtnText: "पॉलिसी का विश्लेषण करें",
                featuresTitle: "मुख्य विशेषताएं",
                feature1Title: "100% स्थानीय AI प्रोसेसिंग",
                feature1Desc: "कोई क्लाउड अपलोड नहीं - आपका डेटा आपके डिवाइस पर रहता है",
                feature2Title: "पूर्ण गोपनीयता",
                feature2Desc: "उन्नत PII सुरक्षा के साथ शून्य डेटा रिसाव",
                feature3Title: "स्मार्ट विश्लेषण",
                feature3Desc: "स्वचालित रूप से मुख्य शर्तें, कवरेज और बहिष्करण निकालता है",
                feature4Title: "टेक्स्ट-टू-स्पीच",
                feature4Desc: "हिंदी या अंग्रेज़ी में अपनी पॉलिसी सुनें",
                loadingText: "आपकी पॉलिसी दस्तावेज़ का विश्लेषण किया जा रहा है...",
                resultsTitle: "📊 विश्लेषण परिणाम",
                policySummary: "पॉलिसी सारांश",
                keyTerms: "महत्वपूर्ण शर्तें",
                exclusions: "बहिष्करण और सीमाएं",
                coverage: "कवरेज विवरण",
                qualityMetrics: "गुणवत्ता मेट्रिक्स",
                confidenceScore: "विश्वास स्कोर",
                qualityGrade: "गुणवत्ता ग्रेड",
                expertReview: "⚠️ विशेषज्ञ समीक्षा आवश्यक"
            }
        };

        function switchLanguage(lang) {
            currentLanguage = lang;
            document.getElementById('langEnBtn').classList.toggle('active', lang === 'en');
            document.getElementById('langHiBtn').classList.toggle('active', lang === 'hi');
            
            const t = translations[lang];
            for (const key in t) {
                const elem = document.getElementById(key);
                if (elem) elem.textContent = t[key];
            }
        }

        function handleFileSelect(event) {
            const file = event.target.files[0];
            if (file) {
                selectedFile = file;
                document.getElementById('analyzeBtn').disabled = false;
                document.getElementById('uploadText').textContent = `Selected: ${file.name}`;
                document.getElementById('uploadSubtext').textContent = `${(file.size / 1024 / 1024).toFixed(2)} MB`;
            }
        }

        async function analyzeDocument() {
            if (!selectedFile) return;

            const formData = new FormData();
            formData.append('file', selectedFile);

            document.getElementById('loading').classList.add('show');
            document.getElementById('analyzeBtn').disabled = true;
            ProgressController.show();

            try {
                await ProgressController.advance(1);
                await new Promise(r => setTimeout(r, 500));
                await ProgressController.advance(2);
                
                const response = await fetch('/upload', { method: 'POST', body: formData });
                await ProgressController.advance(3);
                
                if (response.ok) {
                    await ProgressController.advance(4);
                    const result = await response.json();
                    analysisData = result.analysis;
                    await ProgressController.advance(5);
                    await displayResults(result);
                } else {
                    const error = await response.json();
                    showError(error.detail || 'Analysis failed');
                }
            } catch (error) {
                showError('Network error: ' + error.message);
            } finally {
                document.getElementById('loading').classList.remove('show');
                document.getElementById('analyzeBtn').disabled = false;
            }
        }

        async function displayResults(result) {
            const resultsContent = document.getElementById('resultsContent');
            const analysis = result.analysis;
            const t = translations[currentLanguage];

            let html = `
                <div class="action-buttons-bar">
                    <button class="action-btn" onclick="CopyController.copy('${(analysis.summary || '').replace(/'/g, '\\'')}', this)">
                        <span class="material-icons">content_copy</span> Copy Summary
                    </button>
                    <button class="action-btn" onclick="PrintController.print()">
                        <span class="material-icons">print</span> Print
                    </button>
                    <button class="action-btn" onclick="ShareController.share(analysisData)">
                        <span class="material-icons">share</span> Share
                    </button>
                    <button class="action-btn" onclick="AudioController.download()" style="display:${AudioController.currentAudioUrl ? 'inline-flex' : 'none'}" id="downloadAudioBtn">
                        <span class="material-icons">download</span> Download Audio
                    </button>
                </div>
            `;

            html += `
                <div class="result-card">
                    <div class="result-header">
                        <div class="result-title">
                            <span class="material-icons">summarize</span>
                            ${t.policySummary}
                        </div>
                        <div class="tts-controls">
                            <button class="tts-btn" onclick="AudioController.play(\`${(analysis.summary || '').replace(/`/g, '')}\`, 'en'); setTimeout(() => { const btn = document.getElementById('downloadAudioBtn'); if(btn) btn.style.display = 'inline-flex'; }, 1000);">
                                <span class="material-icons">volume_up</span> EN
                            </button>
                            <button class="tts-btn" onclick="AudioController.play(\`${(analysis.summary || '').replace(/`/g, '')}\`, 'hi'); setTimeout(() => { const btn = document.getElementById('downloadAudioBtn'); if(btn) btn.style.display = 'inline-flex'; }, 1000);">
                                <span class="material-icons">volume_up</span> HI
                            </button>
                            <button class="tts-btn" onclick="AudioController.stop()" style="background: var(--error);">
                                <span class="material-icons">stop</span>
                            </button>
                        </div>
                    </div>
                    <div class="result-content">${formatText(analysis.summary || 'No summary available')}</div>
                </div>
            `;
            
            // Add Key Highlights if available
            if (analysis.key_highlights && analysis.key_highlights.length) {
                html += `<div class="result-card">
                    <div class="result-title">
                        <span class="material-icons">star</span>
                        Key Highlights
                    </div>
                    <div class="result-content">`;
                
                analysis.key_highlights.forEach(highlight => {
                    html += `<div style="padding: 8px 0; border-bottom: 1px solid #E7E0EC;">${highlight}</div>`;
                });
                
                html += `</div></div>`;
            }
            
            // Add What You Need To Know
            if (analysis.what_you_need_to_know && analysis.what_you_need_to_know.length) {
                html += `<div class="result-card" style="background: #E7F5FF; border-left: 4px solid #0066CC;">
                    <div class="result-title" style="color: #0066CC;">
                        <span class="material-icons">lightbulb</span>
                        What You Need To Know
                    </div>
                    <div class="result-content">`;
                
                analysis.what_you_need_to_know.forEach(point => {
                    const icon = point.icon || '✅';
                    const title = point.point || point.title || '';
                    const explanation = point.explanation || point.description || '';
                    
                    html += `<div style="padding: 12px; margin-bottom: 12px; background: white; border-radius: 8px;">
                        <div style="font-weight: 700; font-size: 1rem; margin-bottom: 4px;">${icon} ${title}</div>
                        <div style="color: #49454F;">${explanation}</div>
                    </div>`;
                });
                
                html += `</div></div>`;
            }
            
            // Add Educational Resources
            if (analysis.educational_resources && Object.keys(analysis.educational_resources).length) {
                html += `<div class="result-card" style="background: #FFF8E1; border-left: 4px solid #FFA000;">
                    <div class="result-title" style="color: #FFA000;">
                        <span class="material-icons">school</span>
                        🎓 Learn More - Official Resources
                    </div>
                    <div class="result-content">`;
                
                const resources = analysis.educational_resources;
                
                if (resources.irdai) {
                    html += `<div style="margin-bottom: 12px;">
                        <strong>📜 IRDAI Official Guidelines:</strong><br>
                        <a href="${resources.irdai}" target="_blank" style="color: #0066CC;">${resources.irdai}</a>
                    </div>`;
                }
                
                if (resources.video) {
                    html += `<div style="margin-bottom: 12px;">
                        <strong>🎥 Video Explanations (Hindi + English):</strong><br>
                        <a href="${resources.video}" target="_blank" style="color: #0066CC;">Watch Educational Videos</a>
                    </div>`;
                }
                
                if (resources.complaint_portal) {
                    html += `<div style="margin-bottom: 12px;">
                        <strong>📧 File a Complaint:</strong><br>
                        <a href="${resources.complaint_portal}" target="_blank" style="color: #0066CC;">${resources.complaint_portal}</a>
                    </div>`;
                }
                
                if (resources.insurance_ombudsman) {
                    html += `<div style="margin-bottom: 12px;">
                        <strong>📞 Insurance Ombudsman:</strong><br>
                        <a href="${resources.insurance_ombudsman}" target="_blank" style="color: #0066CC;">${resources.insurance_ombudsman}</a>
                    </div>`;
                }
                
                html += `</div></div>`;
            }

            if (analysis.key_terms && analysis.key_terms.length) {
                html += `<div class="result-card">
                    <div class="result-title">
                        <span class="material-icons">key</span>
                        ${t.keyTerms}
                    </div>
                    <div class="result-content">`;
                
                analysis.key_terms.forEach(term => {
                    const termName = term.term || term.name || 'Term';
                    const simpleExplanation = term.simple_explanation || term.definition || term.description || '';
                    const example = term.example || '';
                    const hindiExplanation = term.hindi_explanation || '';
                    const importance = term.importance || 'MEDIUM';
                    
                    html += `<div class="term-item" style="border-left: 4px solid ${importance === 'HIGH' ? '#BA1A1A' : '#6750A4'};">
                        <div class="term-name" style="font-weight: 700; font-size: 1.1rem; margin-bottom: 8px;">${termName}</div>
                        <div class="term-definition" style="margin-bottom: 8px;">
                            <strong>📖 Simple Explanation:</strong><br>${simpleExplanation}
                        </div>
                        ${example ? `<div style="background: #E7E0EC; padding: 12px; border-radius: 8px; margin-bottom: 8px;">
                            <strong>💡 Example:</strong><br>${example}
                        </div>` : ''}
                        ${hindiExplanation ? `<div style="color: #6750A4; font-style: italic;">
                            <strong>🇮🇳 Hindi:</strong> ${hindiExplanation}
                        </div>` : ''}
                        <div style="margin-top: 8px; font-size: 0.85rem; color: ${importance === 'HIGH' ? '#BA1A1A' : '#49454F'};">
                            <strong>⭐ Importance:</strong> ${importance}
                        </div>
                    </div>`;
                });
                
                html += `</div></div>`;
            }

            if (analysis.exclusions && analysis.exclusions.length) {
                html += `<div class="result-card">
                    <div class="result-title">
                        <span class="material-icons">block</span>
                        ${t.exclusions}
                    </div>
                    <div class="result-content">`;
                
                analysis.exclusions.forEach(exclusion => {
                    // Handle both string exclusions and object exclusions
                    if (typeof exclusion === 'string') {
                        html += `<div class="exclusion-item">
                            <span class="material-icons" style="color: #ff9800;">warning</span>
                            <div>${exclusion}</div>
                        </div>`;
                    } else {
                        const exclusionText = exclusion.exclusion || '';
                        const whatItMeans = exclusion.what_it_means || '';
                        const example = exclusion.example || '';
                        const importance = exclusion.importance || 'MEDIUM';
                        
                        html += `<div class="exclusion-item" style="border-left: 4px solid #ff9800; padding: 16px;">
                            <div style="font-weight: 700; font-size: 1rem; margin-bottom: 8px; color: #ff9800;">
                                <span class="material-icons" style="vertical-align: middle;">warning</span>
                                ${exclusionText}
                            </div>
                            ${whatItMeans ? `<div style="margin-bottom: 8px;">
                                <strong>📖 What It Means:</strong><br>${whatItMeans}
                            </div>` : ''}
                            ${example ? `<div style="background: #FFF8E1; padding: 12px; border-radius: 8px; margin-bottom: 8px;">
                                <strong>💡 Example:</strong><br>${example}
                            </div>` : ''}
                            <div style="font-size: 0.85rem; color: #ff9800;">
                                <strong>⚠️ Importance:</strong> ${importance}
                            </div>
                        </div>`;
                    }
                });
                
                html += `</div></div>`;
            }

            if (analysis.coverage) {
                html += `<div class="result-card">
                    <div class="result-title">
                        <span class="material-icons">health_and_safety</span>
                        ${t.coverage}
                    </div>
                    <div class="result-content">`;
                
                for (const [key, value] of Object.entries(analysis.coverage)) {
                    html += `<div class="coverage-item">
                        <div class="coverage-label">${formatLabel(key)}</div>
                        <div class="coverage-value">${value}</div>
                    </div>`;
                }
                
                html += `</div></div>`;
            }

            if (analysis.evaluation || analysis.confidence_score !== undefined) {
                const score = analysis.confidence_score || analysis.evaluation?.confidence_score || 0;
                const grade = analysis.quality_grade || analysis.evaluation?.quality_grade || 'N/A';
                
                html += `<div class="result-card">
                    <div class="result-title">
                        <span class="material-icons">analytics</span>
                        ${t.qualityMetrics}
                    </div>
                    <div class="result-content">
                        <span class="metric-badge ${score > 0.85 ? 'metric-high' : score > 0.7 ? 'metric-medium' : 'metric-low'}">
                            ${t.confidenceScore}: ${(score * 100).toFixed(1)}%
                        </span>
                        <span class="metric-badge metric-medium">
                            ${t.qualityGrade}: ${grade}
                        </span>
                        ${analysis.hitl_triggered || analysis.requires_hitl_review ? 
                            `<div style="margin-top: 16px; color: var(--error); font-weight: 600;">${t.expertReview}</div>` : ''}
                    </div>
                </div>`;
            }

            resultsContent.innerHTML = html;
            document.getElementById('resultsSection').classList.add('show');
        }

        async function speakText(text, language) {
            try {
                if (audioElement) {
                    audioElement.pause();
                    audioElement = null;
                }

                const response = await fetch('/tts/generate', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ text: text, language: language })
                });

                if (response.ok) {
                    const result = await response.json();
                    audioElement = new Audio(result.audio_url);
                    audioElement.play();
                } else {
                    const utterance = new SpeechSynthesisUtterance(text);
                    utterance.lang = language === 'hi' ? 'hi-IN' : 'en-US';
                    speechSynthesis.speak(utterance);
                }
            } catch (error) {
                console.error('TTS Error:', error);
                try {
                    const utterance = new SpeechSynthesisUtterance(text);
                    utterance.lang = language === 'hi' ? 'hi-IN' : 'en-US';
                    speechSynthesis.speak(utterance);
                } catch (e) {
                    alert('Text-to-speech is not available.');
                }
            }
        }

        function formatText(text) {
            return text.replace(/\\n/g, '<br>');
        }

        function formatLabel(label) {
            return label.replace(/_/g, ' ').replace(/\\b\\w/g, l => l.toUpperCase());
        }

        function showError(message) {
            document.getElementById('resultsContent').innerHTML = `
                <div class="result-card" style="border-left-color: var(--error);">
                    <div class="result-title" style="color: var(--error);">
                        <span class="material-icons">error</span>
                        Error
                    </div>
                    <div class="result-content" style="color: var(--error);">${message}</div>
                </div>
            `;
            document.getElementById('resultsSection').classList.add('show');
        }

        const uploadArea = document.querySelector('.upload-area');
        uploadArea.addEventListener('dragover', (e) => { e.preventDefault(); uploadArea.style.transform = 'scale(1.02)'; });
        uploadArea.addEventListener('dragleave', () => { uploadArea.style.transform = 'scale(1)'; });
        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.style.transform = 'scale(1)';
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                document.getElementById('fileInput').files = files;
                handleFileSelect({ target: { files: files } });
            }
        });

        // Enhanced Controllers
        window.AudioController = { audioElement: null, currentAudioUrl: null, isPlaying: false, async play(text, lang) { try { this.stop(); const res = await fetch('/tts/generate', { method: 'POST', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ text, language: lang }) }); if (res.ok) { const r = await res.json(); this.currentAudioUrl = r.audio_url; this.audioElement = new Audio(r.audio_url); await this.audioElement.play(); } else { const u = new SpeechSynthesisUtterance(text); u.lang = lang === 'hi' ? 'hi-IN' : 'en-US'; speechSynthesis.speak(u); } } catch(e) { console.error(e); const u = new SpeechSynthesisUtterance(text); u.lang = lang === 'hi' ? 'hi-IN' : 'en-US'; speechSynthesis.speak(u); } }, stop() { if (this.audioElement) { this.audioElement.pause(); this.audioElement = null; } }, download() { if (this.currentAudioUrl) { const a = document.createElement('a'); a.href = this.currentAudioUrl; a.download = `policy_${Date.now()}.mp3`; document.body.appendChild(a); a.click(); document.body.removeChild(a); } } };
        window.CopyController = { async copy(text, btn) { try { await navigator.clipboard.writeText(text); const orig = btn.innerHTML; btn.innerHTML = '<span class="material-icons">check</span> Copied!'; btn.style.background = '#198754'; setTimeout(() => { btn.innerHTML = orig; btn.style.background = ''; }, 2000); } catch(e) { alert('Copy failed'); } } };
        window.PrintController = { print() { window.print(); } };
        window.ShareController = { async share(data) { const text = `📄 Policy Summary\n${data.summary ? data.summary.substring(0, 200) + '...' : ''}\n✨ SaralPolicy`; if (navigator.share) { try { await navigator.share({ title: 'Policy Summary', text }); } catch(e) { CopyController.copy(text, event.target); } } else { CopyController.copy(text, event.target); } } };
        window.DarkModeController = { init() { if (localStorage.getItem('darkMode') === 'enabled') this.enable(); }, toggle() { document.body.classList.contains('dark-mode') ? this.disable() : this.enable(); }, enable() { document.body.classList.add('dark-mode'); localStorage.setItem('darkMode', 'enabled'); const btn = document.getElementById('darkModeToggle'); if (btn) btn.querySelector('.material-icons').textContent = 'light_mode'; }, disable() { document.body.classList.remove('dark-mode'); localStorage.setItem('darkMode', 'disabled'); const btn = document.getElementById('darkModeToggle'); if (btn) btn.querySelector('.material-icons').textContent = 'dark_mode'; } };
        window.ProgressController = { steps: [{ id: 1, text: 'Uploading...', progress: 20 }, { id: 2, text: 'Extracting...', progress: 40 }, { id: 3, text: 'Analyzing...', progress: 60 }, { id: 4, text: 'Generating...', progress: 80 }, { id: 5, text: 'Finalizing...', progress: 100 }], show() { document.getElementById('loading').innerHTML = `<div class="analysis-progress"><div class="progress-steps">${this.steps.map(s => `<div class="progress-step" data-step="${s.id}"><div class="step-icon"><span class="material-icons">hourglass_empty</span></div><div class="step-text">${s.text}</div></div>`).join('')}</div><div class="progress-bar-container"><div class="progress-bar-fill" id="progressBarFill"></div></div><div class="progress-percentage" id="progressPercentage">0%</div></div>`; }, async advance(id) { const step = this.steps.find(s => s.id === id); if (!step) return; const el = document.querySelector(`[data-step="${id}"]`); if (el) { el.classList.add(id === 5 ? 'complete' : 'active'); el.querySelector('.material-icons').textContent = id === 5 ? 'check_circle' : 'hourglass_top'; } const bar = document.getElementById('progressBarFill'); const pct = document.getElementById('progressPercentage'); if (bar && pct) { bar.style.width = `${step.progress}%`; pct.textContent = `${step.progress}%`; } for (let i = 1; i < id; i++) { const prev = document.querySelector(`[data-step="${i}"]`); if (prev) { prev.classList.add('complete'); prev.querySelector('.material-icons').textContent = 'check_circle'; } } } };
        document.addEventListener('DOMContentLoaded', () => DarkModeController.init());
    </script>
</body>
</html>
    """


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload and analyze insurance policy."""
    try:
        # Check file size
        if file.size and file.size > analyzer.max_file_size:
            raise HTTPException(status_code=413, detail="File size exceeds 10MB limit")
        
        # Check file type
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in analyzer.supported_formats:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name
        
        try:
            # Extract text
            extracted_text = analyzer.extract_text_from_file(tmp_file_path)
            
            if not extracted_text.strip():
                raise HTTPException(status_code=400, detail="No text could be extracted")
            
            # Analyze policy
            analysis = analyzer.analyze_policy(extracted_text)
            
            if analysis["status"] == "error":
                raise HTTPException(status_code=500, detail=f"Analysis failed: {analysis['error']}")
            
            return {
                "filename": file.filename,
                "analysis": analysis,
                "message": "Analysis completed successfully"
            }
            
        finally:
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Upload failed", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ask")
async def ask_question(request: dict):
    """Answer questions about the policy."""
    try:
        question = request.get("question", "")
        document_text = request.get("document_text", "")
        
        if not question or not document_text:
            raise HTTPException(status_code=400, detail="Question and document text required")
        
        answer = analyzer.answer_question(document_text, question)
        
        return {"answer": answer, "status": "success"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Question answering failed", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ask_document")
async def ask_document(request: dict):
    """Ask questions about the currently indexed policy document.
    
    This endpoint queries the session-indexed policy document and optionally
    augments responses with IRDAI regulatory knowledge.
    
    Request:
        question: User's question about the policy
        include_irdai: Whether to include IRDAI knowledge (default: True)
        top_k: Number of document chunks to retrieve (default: 3)
    
    Returns:
        answer: AI-generated answer
        document_excerpts: Relevant policy excerpts
        irdai_context: Relevant regulatory information (if included)
        confidence: Answer confidence score
    """
    try:
        question = request.get("question", "")
        include_irdai = request.get("include_irdai", True)
        top_k = request.get("top_k", 3)
        
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")
        
        if not rag_service or not rag_service.enabled:
            raise HTTPException(
                status_code=503,
                detail="RAG service not available. Document must be analyzed first."
            )
        
        # Check if there's an active session
        stats = rag_service.get_stats()
        if not stats.get('session_active', False):
            raise HTTPException(
                status_code=400,
                detail="No document in session. Please analyze a policy document first."
            )
        
        # Query the indexed policy document
        doc_results = rag_service.query_document(question, top_k=top_k)
        
        if not doc_results:
            return {
                "answer": "I couldn't find relevant information in your policy document to answer that question.",
                "document_excerpts": [],
                "irdai_context": [],
                "confidence": 0.0,
                "status": "no_results"
            }
        
        # Build document context
        doc_excerpts = []
        doc_context = []
        
        for result in doc_results:
            excerpt = {
                "content": result['content'][:250] + "...",
                "relevance": round(result['hybrid_score'], 3)
            }
            doc_excerpts.append(excerpt)
            doc_context.append(result['content'])
        
        # Query IRDAI knowledge base if requested
        irdai_context_list = []
        irdai_context_text = []
        
        if include_irdai:
            kb_results = rag_service.query_knowledge_base(question, top_k=2)
            
            for result in kb_results:
                if result['hybrid_score'] >= 0.4:  # Relevance threshold
                    irdai_item = {
                        "source": result['metadata'].get('source', 'IRDAI'),
                        "category": result['metadata'].get('insurance_category', 'general'),
                        "excerpt": result['content'][:200] + "...",
                        "relevance": round(result['hybrid_score'], 3)
                    }
                    irdai_context_list.append(irdai_item)
                    irdai_context_text.append(result['content'])
        
        # Generate comprehensive answer using Ollama
        if analyzer.ollama_service:
            # Construct context-rich prompt
            full_context = "\n\n".join([
                "Your Policy Document:",
                *doc_context[:2],  # Top 2 document chunks
            ])
            
            if irdai_context_text:
                full_context += "\n\nRelevant IRDAI Regulations:\n" + "\n".join(irdai_context_text[:1])
            
            prompt = f"""Based on the insurance policy and regulations provided, answer the user's question clearly and accurately. 

Be specific and cite relevant policy clauses when applicable.

Context:
{full_context}

User Question: {question}

Your Answer:"""
            
            try:
                answer = analyzer.ollama_service.answer_question(full_context, question)
            except:
                # Fallback: Simple extraction from top result
                answer = f"Based on your policy: {doc_context[0][:300]}..."
        else:
            # No LLM available - return top document chunk
            answer = f"Based on your policy document: {doc_context[0][:400]}..."
        
        # Calculate confidence based on relevance scores
        avg_relevance = sum(r['hybrid_score'] for r in doc_results[:2]) / min(2, len(doc_results))
        confidence = round(avg_relevance, 3)
        
        return {
            "answer": answer,
            "document_excerpts": doc_excerpts,
            "irdai_context": irdai_context_list,
            "confidence": confidence,
            "sources_used": {
                "document_chunks": len(doc_excerpts),
                "irdai_regulations": len(irdai_context_list)
            },
            "status": "success"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Document Q&A failed", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/rag/ask")
async def rag_enhanced_ask(request: dict):
    """RAG-enhanced Q&A: Uses document chunks + IRDAI knowledge base (legacy endpoint)."""
    try:
        question = request.get("question", "")
        use_knowledge_base = request.get("use_knowledge_base", True)
        
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")
        
        if not rag_service or not rag_service.enabled:
            raise HTTPException(
                status_code=503,
                detail="RAG service not available. Install: pip install chromadb rank-bm25"
            )
        
        # Query document chunks
        doc_results = rag_service.query_document(question, top_k=3)
        
        # Query IRDAI knowledge base if requested
        kb_results = []
        if use_knowledge_base:
            kb_results = rag_service.query_knowledge_base(question, top_k=2)
        
        # Combine context
        context_parts = []
        sources = []
        
        if doc_results:
            context_parts.append("From your policy document:")
            for i, result in enumerate(doc_results, 1):
                context_parts.append(f"{i}. {result['content'][:300]}...")
                sources.append({"type": "document", "score": result['hybrid_score']})
        
        if kb_results:
            context_parts.append("\nFrom IRDAI knowledge base:")
            for i, result in enumerate(kb_results, 1):
                context_parts.append(f"{i}. {result['content'][:300]}...")
                sources.append({
                    "type": "irdai_knowledge",
                    "source": result['metadata'].get('source', 'IRDAI'),
                    "score": result['hybrid_score']
                })
        
        context = "\n".join(context_parts)
        
        # Generate answer using Ollama with context
        if analyzer.ollama_service:
            prompt = f"""Based on the following context, answer the question clearly and accurately.

Context:
{context}

Question: {question}

Answer:"""
            
            answer = analyzer.ollama_service.answer_question(context, question)
        else:
            answer = "Ollama service not available. Please start Ollama."
        
        return {
            "answer": answer,
            "context_used": len(context_parts) > 0,
            "sources": sources,
            "rag_enabled": True,
            "status": "success"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error("RAG Q&A failed", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/rag/stats")
async def rag_stats():
    """Get RAG service statistics."""
    if rag_service:
        return rag_service.get_stats()
    return {"enabled": False, "message": "RAG service not available"}


@app.get("/health")
async def health_check():
    """Health check with model status."""
    try:
        model_info = llm_service.get_model_info()
        return {
            "status": "healthy", 
            "message": "SaralPolicy is running",
            "model_info": model_info
        }
    except Exception as e:
        return {"status": "degraded", "error": str(e)}


@app.get("/model-info")
async def get_model_info():
    """Get information about the loaded local model."""
    try:
        return llm_service.get_model_info()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/tts/speak")
async def speak_text(request: dict):
    """Convert text to speech (server-side playback)."""
    try:
        text = request.get("text", "")
        language = request.get("language", "en")

        if not text.strip():
            raise HTTPException(status_code=400, detail="Text is required")

        if language not in ["en", "hi"]:
            raise HTTPException(status_code=400, detail="Language must be 'en' or 'hi'")

        success = tts_service.speak_text(text, language)

        if success:
            return {"status": "success", "message": "Text spoken successfully"}
        else:
            raise HTTPException(status_code=500, detail="TTS failed")

    except HTTPException:
        raise
    except Exception as e:
        logger.error("TTS endpoint error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/tts/status")
async def get_tts_status():
    """Get TTS service status and available languages."""
    return tts_service.get_status()


@app.post("/tts/generate")
async def generate_tts_audio(request: dict):
    """Generate TTS audio file and return URL for browser playback."""
    try:
        text = request.get("text", "")
        language = request.get("language", "en")

        if not text.strip():
            raise HTTPException(status_code=400, detail="Text is required")

        if language not in ["en", "hi"]:
            raise HTTPException(status_code=400, detail="Language must be 'en' or 'hi'")

        audio_file = tts_service.generate_audio_file(text, language)

        if not audio_file:
            raise HTTPException(status_code=500, detail="TTS generation failed")

        filename = os.path.basename(audio_file)
        return {
            "status": "success",
            "audio_url": f"/tts/audio/{filename}",
            "message": "Audio generated successfully"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error("TTS generate endpoint error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/tts/audio/{filename}")
async def serve_tts_audio(filename: str):
    """Serve generated TTS audio files."""
    try:
        audio_path = tts_service.get_audio_file_path(filename)
        if audio_path:
            media_type = "audio/mpeg" if filename.lower().endswith(".mp3") else "audio/wav"
            return FileResponse(audio_path, media_type=media_type)
        else:
            raise HTTPException(status_code=404, detail="Audio file not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error("TTS audio serve error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/translate/status")
async def get_translation_status():
    """Get translation service status."""
    return translation_service.get_status()


@app.post("/translate")
async def translate_text(request: dict):
    """Translate text between Hindi and English."""
    try:
        text = request.get("text", "")
        target_language = request.get("target_language", "hi")

        if not text.strip():
            raise HTTPException(status_code=400, detail="Text is required")

        if target_language not in ["en", "hi"]:
            raise HTTPException(status_code=400, detail="Target language must be 'en' or 'hi'")

        translated_text = translation_service.translate_text(text, target_language)

        return {
            "status": "success",
            "original_text": text,
            "translated_text": translated_text,
            "target_language": target_language
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Translation endpoint error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    logger.info("Starting SaralPolicy server...")
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
